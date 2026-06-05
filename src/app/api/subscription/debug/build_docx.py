from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles helper ─────────────────────────────────────────────────────────────
def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=True, color=(31, 73, 125))
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        set_font(run, size=13, bold=True, color=(21, 96, 130))
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        set_font(run, size=11, bold=True, color=(54, 95, 145))
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
    # bottom border for h1
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F497D')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def add_body(doc, text, bold=False, italic=False, size=10.5, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p

def add_bullet(doc, text, level=0, marker="•"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(0.6 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(1)
    run = p.add_run(f"{marker}  {text}")
    set_font(run, size=10.5)
    return p

def add_code_block(doc, text):
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'F0F0F0')
        p._p.get_or_add_pPr().append(shading)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None, header_color="1F497D"):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        shade_cell(cell, header_color)
        run = cell.paragraphs[0].runs[0]
        set_font(run, size=10, bold=True, color=(255, 255, 255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri+1]
        bg = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = val
            shade_cell(cell, bg)
            run = cell.paragraphs[0].runs[0]
            set_font(run, size=10)
    # col widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacing after table
    return table

def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_warning(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("⚠  " + text)
    set_font(run, size=10, bold=True, color=(180, 95, 0))
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'FFF3CD')
    p._p.get_or_add_pPr().append(shading)

# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run("BRIEFING HANDOVER")
set_font(r, size=26, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SILVERIUM PROJECT")
set_font(r, size=22, bold=True, color=(21, 96, 130))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
r = p.add_run("Tanggal: 5 Juni 2026  |  Jaringan: Polygon Mainnet  |  Domain: wallet.silverium.id")
set_font(r, size=11, italic=True, color=(100, 100, 100))

doc.add_paragraph()
add_hr(doc)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — Gambaran Umum
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Gambaran Umum Project")
add_body(doc, "Silverium adalah platform token RWA (Real World Asset) di Polygon Mainnet. Fitur utama:")
add_bullet(doc, "Swap SILV ↔ USDT langsung dari dashboard")
add_bullet(doc, "Stake SILV dengan pilihan lock 3 bulan (8% APY), 6 bulan (12%), atau 12 bulan (18%) dan mendapat reward USDT")
add_bullet(doc, "Admin mengelola lot emas, PnL, dividen, treasury, dan monitoring posisi staking via panel admin")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — Arsitektur
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Arsitektur Sistem")
add_code_block(doc, """User Browser
    │
    ▼
Nginx (HTTPS)
    ├── wallet.silverium.id    → port 5000 → Next.js (Frontend)
    └── apicoin.silverium.id   → port 4000 → Express API (Backend)
                                                    │
                                                    ▼
                                             PostgreSQL (Prisma)

Frontend ←→ Polygon Mainnet (RPC: Alchemy)
Frontend ←→ SushiSwap V3 Pool (swap & harga SILV)
Frontend ←→ Staking Contract (stake / unstake / claim)
Admin Panel ←→ Express API ←→ Database""")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — Stack
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Stack Teknologi")
add_table(doc,
    ["Layer", "Teknologi"],
    [
        ["Frontend", "Next.js 14 App Router + Tailwind CSS"],
        ["Blockchain Library", "Viem + Wagmi v2"],
        ["Wallet Modal", "Web3Modal (WalletConnect)"],
        ["Backend API", "Express.js + TypeScript"],
        ["Database ORM", "Prisma + PostgreSQL"],
        ["Process Manager", "Dijalankan manual via terminal (harus pindah ke PM2)"],
        ["Web Server", "Nginx + Let's Encrypt SSL"],
    ],
    col_widths=[5, 11]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — Struktur File
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. Struktur File — Penjelasan Lengkap")

add_heading(doc, "4.1 Frontend (silverium-web/src/)", level=2)

add_heading(doc, "app/ — Halaman Utama", level=3)
add_table(doc,
    ["File", "Fungsi"],
    [
        ["app/page.tsx", "FILE UTAMA — semua UI user ada di sini (dashboard, swap, staking, history)"],
        ["app/layout.tsx", "Root layout, hanya membungkus WalletProvider"],
        ["app/globals.css", "Global CSS styles"],
        ["app/price/route.ts", "Next.js API Route GET /price — fetch harga POL & IDR dari CoinGecko"],
        ["app/admin/layout-client.tsx", "Guard: cek admin_token, redirect ke /login kalau tidak ada token"],
        ["app/admin/login/page.tsx", "Form login admin (hit POST /admin/login ke API)"],
        ["app/admin/page.tsx", "Dashboard ringkasan admin"],
        ["app/admin/lots/page.tsx", "Manajemen lot emas"],
        ["app/admin/pnl/page.tsx", "Input profit & loss per periode"],
        ["app/admin/dividends/page.tsx", "Distribusi dividen ke holder"],
        ["app/admin/staking/page.tsx", "Monitor posisi staking semua user"],
        ["app/admin/reserve/page.tsx", "Monitor treasury USDT"],
        ["app/admin/profile/page.tsx", "Profil & info akun admin"],
    ],
    col_widths=[5.5, 10.5]
)

add_heading(doc, "components/ — Komponen UI", level=3)
add_table(doc,
    ["File", "Status", "Fungsi"],
    [
        ["components/Sidebar.tsx", "✅ Aktif", "Navigasi kiri — dipakai layout admin dan web3"],
        ["components/Topbar.tsx", "✅ Aktif", "Header atas — berisi tombol logout admin (pakai auth.js)"],
        ["components/web3/connectwallet.tsx", "⚠ Dipertahankan", "Komponen connect wallet legacy — dipertahankan"],
        ["components/web3/eb3Provider.tsx", "⚠ Dipertahankan", "Web3 provider alternatif legacy — dipertahankan"],
        ["components/ui/*", "✅ Aktif", "Komponen UI: button, card, input, dialog, table, tabs, sonner, dll"],
    ],
    col_widths=[5.5, 3, 7.5]
)

add_heading(doc, "providers/ & lib/", level=3)
add_table(doc,
    ["File", "Status", "Fungsi"],
    [
        ["providers/wallet.tsx", "✅ INTI", "WagmiConfig + QueryClient + WalletProvider — JANGAN diubah sembarangan"],
        ["lib/api.ts", "✅ Aktif", "Helper fetch ke API (get, post, patch) + baca/hapus admin_token"],
        ["lib/auth.js", "✅ Aktif", "adminLogout() — dipakai Topbar.tsx untuk tombol logout"],
        ["lib/chain.ts", "✅ Aktif", "currentChain config — dipakai admin/profile/page.tsx"],
        ["lib/utils.ts", "✅ Aktif", "cn() untuk className merge — dipakai semua komponen UI"],
    ],
    col_widths=[5, 3, 8]
)

add_heading(doc, "4.2 Backend API (api/src/)", level=2)

add_heading(doc, "routes/ — Endpoint API", level=3)
add_table(doc,
    ["Route", "File", "Fungsi"],
    [
        ["GET/POST /history", "routes/history.ts", "Riwayat transaksi user (swap, stake, claim)"],
        ["/admin/*", "routes/admin.ts", "Login admin, CRUD, treasury management"],
        ["/lots", "routes/lots.ts", "Manajemen lot emas"],
        ["/pnl", "routes/pnl.ts", "Profit & loss per periode"],
        ["/dividends", "routes/dividends.ts", "Distribusi dividen"],
        ["/silv", "routes/silv.ts", "Info token SILV on-chain"],
        ["/snapshot", "routes/snapshot.ts", "Snapshot balance holder"],
        ["/swap", "routes/swap.ts", "Endpoint swap (saat ini tidak dipakai frontend)"],
    ],
    col_widths=[4, 5, 7]
)

add_heading(doc, "services/ — Service Layer", level=3)
add_table(doc,
    ["File", "Dipakai Oleh", "Fungsi"],
    [
        ["historyIndexer.ts", "routes/history.ts", "Fetch & simpan riwayat transaksi ke database"],
        ["onchain.ts", "routes/dividends.ts", "Baca data on-chain (balance, event)"],
        ["buildCalldata.ts", "routes/admin.ts", "Encode calldata untuk transaksi admin"],
        ["validate.ts", "routes/admin.ts", "Validasi input (BPS, address, batch transfer)"],
    ],
    col_widths=[4.5, 5, 6.5]
)

add_heading(doc, "4.3 Database — Model Prisma", level=2)
add_table(doc,
    ["Model", "Fungsi"],
    [
        ["Lot", "Data lot emas (serial, berat, kemurnian)"],
        ["Lease", "Data penyewaan lot"],
        ["ProductionBatch", "Produksi dari lot emas"],
        ["Sale", "Data penjualan"],
        ["PnlPeriod", "Periode profit & loss"],
        ["Dividend", "Distribusi dividen per periode"],
        ["DividendAllocation", "Alokasi dividen per holder wallet"],
        ["WalletBalance", "Balance SILV per address"],
        ["Snapshot", "Snapshot balance untuk perhitungan dividen"],
        ["UserHistory", "Riwayat transaksi user (swap, stake, claim)"],
        ["OnchainEvent", "Event on-chain yang di-index"],
        ["Admin", "Akun admin (username + hashed password)"],
    ],
    col_widths=[5, 11]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — Konfigurasi
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. File Konfigurasi Penting")

add_heading(doc, "5.1 .env.local (Frontend)", level=2)
add_body(doc, "Lokasi di VPS: /var/www/silverium/silverium-web/.env.local", italic=True)
add_table(doc,
    ["Variable", "Keterangan"],
    [
        ["NEXT_PUBLIC_SILV_ADDRESS", "Alamat token SILV di Polygon"],
        ["NEXT_PUBLIC_USDT_ADDRESS", "Alamat USDT Polygon"],
        ["NEXT_PUBLIC_STAKING_V3", "Alamat staking contract production"],
        ["NEXT_PUBLIC_SILV_USDT_V3_POOL", "Alamat pool SushiSwap V3 SILV/USDT"],
        ["NEXT_PUBLIC_V3_ROUTER", "SushiSwap V3 SwapRouter — dipakai handleSwap()"],
        ["NEXT_PUBLIC_SUSHI_V3_QUOTER", "QuoterV2 — untuk estimasi output swap"],
        ["NEXT_PUBLIC_API_BASE", "https://apicoin.silverium.id"],
        ["NEXT_PUBLIC_RPC_URL_POLYGON", "Alchemy RPC — jangan ganti sembarangan"],
        ["NEXT_PUBLIC_WC_PROJECT_ID", "WalletConnect Project ID"],
    ],
    col_widths=[7, 9]
)

add_heading(doc, "5.2 next.config.js — Hal yang Perlu Diketahui", level=2)
add_bullet(doc, "output: 'standalone' — sudah siap untuk production build")
add_bullet(doc, "allowedDevOrigins — daftar origin yang boleh akses dev server. Bisa dihapus setelah pindah ke production build")
add_bullet(doc, "webpack alias — memblok porto dan @metamask/sdk karena ada bug dependency konflik. JANGAN dihapus")
add_bullet(doc, "rewrites /api/* → localhost:3001 — tidak dipakai frontend saat ini, API dipanggil langsung ke apicoin.silverium.id")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — Server (masalah kritis)
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. Cara Server Dijalankan — Masalah Kritis")

add_warning(doc, "MASALAH KRITIS: Web dan API keduanya berjalan dalam mode DEVELOPMENT, bukan production. Next dev memakai ~3GB RAM dan performa jauh lebih lambat dari production build.")

add_heading(doc, "Yang sedang berjalan sekarang (SALAH):", level=3)
add_code_block(doc, "next dev -p 5000        ← Frontend (DEV MODE)\nts-node src/index.ts    ← API (DEV MODE)")

add_heading(doc, "Yang seharusnya dijalankan (production):", level=3)
add_code_block(doc, """# Frontend
cd /var/www/silverium/silverium-web
npm run build                   # build production
next start -p 5000              # jalankan hasil build

# API
cd /var/www/silverium/api
npx tsc                         # compile TypeScript
node dist/index.js              # jalankan compiled output""")

add_heading(doc, "Kelola dengan PM2 agar tidak mati saat terminal ditutup:", level=3)
add_code_block(doc, """pm2 start "next start -p 5000" --name silverium-web
pm2 start "node dist/index.js" --name silverium-api
pm2 save
pm2 startup""")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — Perubahan Kode
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. Perubahan Kode yang Sudah Dilakukan")

add_heading(doc, "7.1 page.tsx — 4 Perubahan", level=2)

add_heading(doc, "Perubahan 1: handleSwap() — Bypass SushiSwap API", level=3)
add_table(doc,
    ["", "Keterangan"],
    [
        ["Sebelum", "Swap lewat api.sushi.com (API eksternal) → generate calldata → kirim transaksi"],
        ["Sesudah", "Swap langsung ke SUSHI_V3_SWAP_ROUTER via writeContractAsync + exactInputSingle"],
        ["Kenapa", "SushiSwap API menolak (HTTP 422) karena pool likuiditas terlalu kecil ($3.58), tidak memenuhi minimum threshold API mereka"],
        ["Efek", "Swap tidak bergantung API eksternal apapun, langsung ke smart contract — lebih andal"],
    ],
    col_widths=[3, 13]
)

add_heading(doc, "Perubahan 2, 3, 4: handleStake / handleUnstake / handleClaimStakingReward — Cek Status Receipt", level=3)
add_table(doc,
    ["", "Keterangan"],
    [
        ["Sebelum", "Tunggu transaksi selesai → langsung toast sukses (tanpa cek apakah benar sukses di blockchain)"],
        ["Sesudah", "Cek receipt.status !== 'success' → kalau gagal, throw error → toast merah muncul"],
        ["Kenapa", "Ada kasus transaksi dikirim tapi di-revert oleh contract (gagal di blockchain), namun UI tetap tampil sukses — menyesatkan user"],
        ["Efek", "Tidak ada lagi 'sukses palsu'. Semua kegagalan transaksi akan tampil sebagai notifikasi error"],
    ],
    col_widths=[3, 13]
)

add_heading(doc, "7.2 File Dihapus dari VPS", level=2)
add_body(doc, "Backup tersimpan di: /var/www/silverium/_backup_20260605_102331.tar.gz", italic=True)
add_table(doc,
    ["File / Folder", "Alasan Dihapus"],
    [
        ["app/0 dan app/0n", "File kosong tanpa konten"],
        ["app/page1.tsx", "Backup lama dari page.tsx — sudah tidak relevan"],
        ["app/client-layout.tsx", "Tidak diimport oleh file manapun"],
        ["app/(web3)/ — seluruh folder", "Route group lama, tidak ada link dari UI, tidak dipakai"],
        ["lib/token.ts", "Tidak diimport oleh file manapun"],
        ["lib/format.ts", "Tidak diimport oleh file manapun"],
        ["lib/safe.ts", "Tidak diimport oleh file manapun"],
        ["api/src/services/onchainSync.ts", "Tidak diimport oleh route manapun"],
        ["api/src/src/ (folder duplikat)", "Folder nested hasil copy yang salah"],
    ],
    col_widths=[6, 10]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — BOLEH vs TIDAK BOLEH diubah
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. Yang BOLEH dan TIDAK BOLEH Diubah")

add_heading(doc, "✅ BOLEH Diubah", level=2)
add_table(doc,
    ["File / Bagian", "Yang Boleh Diubah"],
    [
        ["page.tsx", "Tampilan UI, teks, warna, layout section dashboard"],
        ["app/admin/*", "Tambah fitur admin baru, perbaiki tampilan tabel"],
        ["api/src/routes/*", "Tambah endpoint baru, perbaiki logika bisnis"],
        [".env.local", "Ganti RPC URL, alamat contract jika deploy ulang"],
        ["next.config.js — allowedDevOrigins", "Hapus setelah pindah ke production build"],
    ],
    col_widths=[5, 11]
)

add_heading(doc, "❌ TIDAK BOLEH Diubah Sembarangan", level=2)
add_table(doc,
    ["File / Bagian", "Kenapa"],
    [
        ["providers/wallet.tsx", "Inti koneksi wallet seluruh app. Ubah salah → semua fitur blockchain mati"],
        ["SUSHI_V3_SWAP_ROUTER di page.tsx", "Router yang terbukti bekerja untuk pool SILV/USDT. Jangan ganti ke Uniswap Router (beda factory hash, akan gagal)"],
        ["Konstanta SLIPPAGE_BPS = 500", "5% slippage sudah disesuaikan untuk pool kecil. Kalau dikecilkan, swap akan sering gagal karena MinimalOutputViolation"],
        ["stakingV3Abi di page.tsx", "ABI harus persis sesuai contract yang di-deploy. Jika contract di-redeploy, ABI wajib diupdate sesuai contract baru"],
        ["api/prisma/schema.prisma", "Jangan ubah tanpa jalankan migration (npx prisma migrate). Ubah schema tanpa migrate akan crash API"],
        ["Nginx CORS config apicoin.silverium.id", "Sudah dikonfigurasi agar hanya wallet.silverium.id yang bisa akses API. Jangan ubah origin sembarangan"],
        ["webpack alias di next.config.js", "Memblok dependency bermasalah (porto, @metamask/sdk). Jangan dihapus meski terlihat tidak perlu"],
    ],
    col_widths=[5, 11]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — Alur Teknis
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "9. Alur Teknis Fitur Utama")

add_heading(doc, "9.1 Alur Swap", level=2)
add_code_block(doc, """User input amount
    → quoteViaSimulation()
        → Coba V3 Quoter contract (fee: 3000, 10000, 500, 100)
        → Fallback: hitung manual dari sqrtPriceX96 yang di-cache
    → Tampil estimasi output + min received (dengan 5% slippage)
    → handleSwap():
        1. Cek allowance token ke SUSHI_V3_SWAP_ROUTER
        2. Kalau kurang → approve (USDT perlu reset ke 0 dulu sebelum approve baru)
        3. writeContractAsync(exactInputSingle) ke SUSHI_V3_SWAP_ROUTER
        4. waitForTransactionReceipt → cek receipt.status
        5. POST ke /history untuk simpan riwayat""")

add_heading(doc, "9.2 Alur Stake", level=2)
add_code_block(doc, """User input amount + pilih lock period (3m / 6m / 1y)
    → Cek stakingAllowance ke STAKING_V3
    → Kalau kurang → handleApproveStakeToken() (approve maxUint256)
    → handleStake():
        1. parseUnits(amount, silvDecimals)
        2. Map lock period ke detik: 3m=90 hari, 6m=180 hari, 1y=365 hari
        3. writeContractAsync(stake, [amount, lockPeriodInSeconds])
        4. waitForTransactionReceipt → cek receipt.status
        5. POST ke /history untuk simpan riwayat""")

add_heading(doc, "9.3 Alur Harga SILV (On-Chain)", level=2)
add_code_block(doc, """slot0 poller setiap 30 detik
    → readContract(pool.slot0) → dapat sqrtPriceX96
    → priceFromSqrtX96() → hitung harga SILV dalam USDT
    → setSilvPriceUsd() → tampil di UI dashboard
    → Dipakai juga sebagai fallback quote kalau Quoter contract gagal""")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — Alamat Contract
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "10. Alamat Contract")
add_table(doc,
    ["Nama", "Alamat"],
    [
        ["Token SILV", "0xa2f471e8701CeA11f71D2ADdD37B317BaC266028"],
        ["Staking Production", "0x8dD8bfF9aea93D193F1Adf5426309b33F6c852E4"],
        ["Staking Test (lock 5 menit)", "0xfCd02688BdBE9D3D07A7DB18428b9AFfBEaa3f98"],
        ["Pool SILV/USDT V3", "0x520c0af6df9a66ea63aed1138cfec251938f6f7c"],
        ["Treasury / Safe Wallet", "0x83A74f4F7708140C620BCdE9018BE3b7Ea3635b7"],
        ["USDT Polygon", "0xc2132D05D31c914a87C6611C10748AEb04B58e8F"],
        ["SushiSwap V3 Router", "0x8a21F6768C1f8075727fDf3d53C8BcDBA6A9B40f"],
    ],
    col_widths=[6, 10]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — Task Berikutnya
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "11. Yang Harus Dikerjakan Developer Berikutnya")
add_table(doc,
    ["Prioritas", "Task", "Detail"],
    [
        ["🔴 Segera", "Pindah ke production build", "Stop next dev, jalankan npm run build && next start, kelola via PM2"],
        ["🔴 Segera", "Compile API ke JavaScript", "Jalankan npx tsc lalu node dist/index.js, kelola via PM2"],
        ["🔴 Segera", "Isi Treasury USDT", "Saat ini hanya $0.05 — user tidak bisa claim reward"],
        ["🔴 Segera", "Tambah Liquidity Pool", "Saat ini hanya $3.58 — swap amount besar akan slippage ekstrem. Minimal $50-100 USDT"],
        ["🟡 Penting", "Tambah signer kedua Safe Wallet", "Supaya akses tidak bergantung pada 1 orang saja"],
        ["🟡 Penting", "Transfer ownership contract ke client", "Ikuti dokumen 5 Fase di Google Drive — Fase 3 IRREVERSIBLE, tidak bisa dibatalkan"],
        ["🟢 Tunggu", "PolygonScan token update", "Sudah disubmit — estimasi 3-5 hari kerja"],
        ["🟢 Tunggu", "GeckoTerminal listing", "Sudah disubmit — estimasi 5 hari kerja"],
        ["🔐 Keamanan", "Ganti password VPS", "Password root dipakai dalam sesi handover ini — wajib diganti segera"],
    ],
    col_widths=[3, 5, 8]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — Urutan Eksekusi
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "12. Urutan Eksekusi Handover (Jangan Dibalik)")
add_table(doc,
    ["Fase", "Aksi", "Risk"],
    [
        ["Fase 1 — Keamanan", "Buat wallet baru, revoke approve wallet lama yang terkompromisi, tambah signer Safe", "Rendah"],
        ["Fase 2 — Operasional", "Isi treasury USDT, approve USDT ke staking contract, tambah liquidity pool", "Rendah"],
        ["Fase 3 — Handover", "Transfer ownership contract ke client, serahkan akses Safe Wallet", "⚠ IRREVERSIBLE"],
        ["Fase 4 — Monitoring", "Pantau pool, posisi user, transaksi, aktivitas swap", "Nol"],
        ["Fase 5 — Darurat", "Emergency withdraw, nonaktifkan staking, ganti treasury jika ada bug", "Tinggi"],
    ],
    col_widths=[4.5, 9, 2.5]
)

add_warning(doc, "Fase 3 tidak bisa di-undo. Pastikan Fase 1 dan 2 sudah berjalan normal minimal beberapa hari sebelum eksekusi transfer ownership.")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 13 — Dokumentasi
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "13. Dokumentasi Operasional")
add_body(doc, "19 dokumen panduan operasional tersimpan di Google Drive folder Dokumentasi, mencakup:")
docs_list = [
    "Cara Menambah Liquidity Pool",
    "Cara Mengisi Saldo Treasury USDT",
    "Cara Mengubah APY",
    "Cara Memantau Kesehatan Pool Staking",
    "Cara Memantau Posisi dan Reward User",
    "Cara Approve USDT Treasury ke Staking Contract",
    "Cara Mengganti Treasury Address",
    "Cara Menambah Signer Safe Wallet",
    "Cara Transfer Ownership Contract",
    "Prosedur Jika Signer Safe Wallet Hilang",
    "Cara Mengaktifkan Emergency Withdraw",
    "Prosedur Jika Treasury Kehabisan USDT",
    "Cara Menonaktifkan Staking Jika Ada Bug",
    "Cara Revoke Approve Wallet yang Terkompromisi",
    "Cara Membuat Wallet Baru yang Aman",
    "Cara Cek Riwayat Transaksi Token SILV",
    "Cara Memperkirakan Kebutuhan Treasury per Periode",
    "Cara Memantau Aktivitas Swap di Pool SushiSwap",
    "Cara Remove Liquidity",
]
for i, d in enumerate(docs_list, 1):
    add_bullet(doc, f"{i:02d}. {d}")

add_body(doc, "")
add_body(doc, "Dokumen Urutan Teknis Pemindahan ke Client (5 Fase) wajib dibaca sebelum melakukan transfer ownership apapun.", bold=True)

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_hr(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Dokumen ini mencakup kondisi project per 5 Juni 2026. Untuk detail operasional, rujuk 19 dokumen di Google Drive.")
set_font(r, size=9, italic=True, color=(120, 120, 120))

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/Users/akm/Documents/Silverium/HANDOVER_SILVERIUM_2026.docx"
doc.save(out)
print(f"Saved: {out}")
